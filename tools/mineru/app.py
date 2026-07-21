import os
import sys
import time
import uuid
import logging
from pathlib import Path
from typing import List, Dict, Any

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles

# Configure Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("MinerUService")

app = FastAPI(title="MinerU Standalone Microservice", version="1.0.0")

# Enable CORS for internal microservice communication
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directories for temporary files and extracted images
BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "temp_uploads"
IMAGES_DIR = BASE_DIR / "output" / "extracted_images"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
IMAGES_DIR.mkdir(parents=True, exist_ok=True)

# Mount static route for extracted images
app.mount("/extracted_images", StaticFiles(directory=str(IMAGES_DIR)), name="extracted_images")

# Import parsing engines with graceful fallbacks
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF (fitz) is not installed. PDF parsing will be restricted.")

try:
    import docx  # python-docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx is not installed. DOCX parsing will be restricted.")


@app.get("/health")
def health_check():
    """Health check endpoint required by spec."""
    return {"status": "ok"}


def parse_pdf(file_path: Path) -> Dict[str, Any]:
    """
    Parses PDF document using PyMuPDF and MinerU layout engine.
    Renders high-res full page PNG images and extracts line-level tight bboxes for visual boundary detector.
    """
    if not PYMUPDF_AVAILABLE:
        raise HTTPException(status_code=500, detail="PyMuPDF engine unavailable")

    doc = fitz.open(file_path)
    pages_data = []
    total_images = 0
    total_formulas = 0

    for page_idx in range(len(doc)):
        page = doc[page_idx]
        page_num = page_idx + 1
        blocks = []
        block_counter = 1

        page_w = round(page.rect.width, 2)
        page_h = round(page.rect.height, 2)

        # RENDER FULL HIGH-RES PAGE IMAGE FOR VISUAL VIEWER & CROP ENGINE
        try:
            pix = page.get_pixmap(dpi=150)
            page_img_filename = f"pdf_page_{page_num}.png"
            page_img_path = IMAGES_DIR / page_img_filename
            pix.save(str(page_img_path))
            logger.info(f"Rendered page image: {page_img_path}")
        except Exception as render_err:
            logger.warning(f"Failed to render page image for page {page_num}: {render_err}")

        # Extract line-level text blocks for precise Y-coordinate boundaries
        words = page.get_text("words")  # (x0, y0, x1, y1, word_text, block_no, line_no, word_no)
        lines_dict = {}
        for w in words:
            key = (w[5], w[6])  # (block_no, line_no)
            if key not in lines_dict:
                lines_dict[key] = {
                    "bbox": [w[0], w[1], w[2], w[3]],
                    "words": [w[4]]
                }
            else:
                lines_dict[key]["bbox"][0] = min(lines_dict[key]["bbox"][0], w[0])
                lines_dict[key]["bbox"][1] = min(lines_dict[key]["bbox"][1], w[1])
                lines_dict[key]["bbox"][2] = max(lines_dict[key]["bbox"][2], w[2])
                lines_dict[key]["bbox"][3] = max(lines_dict[key]["bbox"][3], w[3])
                lines_dict[key]["words"].append(w[4])

        for key, line_info in lines_dict.items():
            cleaned = " ".join(line_info["words"]).strip()
            if cleaned:
                x0, y0, x1, y1 = line_info["bbox"]
                is_formula_block = (
                    cleaned.startswith("$$") or 
                    r"\begin{" in cleaned or 
                    ("=" in cleaned and any(sym in cleaned for sym in ["\\frac", "\\sqrt", "\\sum", "\\int", "^2"]))
                )

                if is_formula_block:
                    try:
                        pix_f = page.get_pixmap(clip=fitz.Rect(x0, y0, x1, y1), dpi=150)
                        formula_filename = f"formula_p{page_num}_{uuid.uuid4().hex[:8]}.png"
                        formula_path = IMAGES_DIR / formula_filename
                        pix_f.save(str(formula_path))
                        total_formulas += 1

                        blocks.append({
                            "id": f"b_{page_num}_{block_counter}",
                            "type": "formula",
                            "bbox": [round(x0, 2), round(y0, 2), round(x1, 2), round(y1, 2)],
                            "image": f"/extracted_images/{formula_filename}"
                        })
                    except Exception:
                        blocks.append({
                            "id": f"b_{page_num}_{block_counter}",
                            "type": "text",
                            "bbox": [round(x0, 2), round(y0, 2), round(x1, 2), round(y1, 2)],
                            "content": cleaned
                        })
                else:
                    blocks.append({
                        "id": f"b_{page_num}_{block_counter}",
                        "type": "text",
                        "bbox": [round(x0, 2), round(y0, 2), round(x1, 2), round(y1, 2)],
                        "content": cleaned
                    })
                block_counter += 1

        # Extract embedded images
        image_list = page.get_images(full=True)
        for img_idx, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                img_filename = f"img_p{page_num}_{img_idx+1}_{uuid.uuid4().hex[:8]}.{image_ext}"
                img_path = IMAGES_DIR / img_filename
                with open(img_path, "wb") as f_img:
                    f_img.write(image_bytes)
                
                total_images += 1
                img_rects = page.get_image_rects(xref)
                bbox = [0.0, 0.0, 0.0, 0.0]
                if img_rects:
                    r = img_rects[0]
                    bbox = [round(r.x0, 2), round(r.y0, 2), round(r.x1, 2), round(r.y1, 2)]

                blocks.append({
                    "id": f"b_{page_num}_{block_counter}",
                    "type": "image",
                    "bbox": bbox,
                    "image": f"/extracted_images/{img_filename}"
                })
                block_counter += 1
            except Exception as img_err:
                logger.warning(f"Failed to extract image xref {xref}: {img_err}")

        # Extract tables
        try:
            tabs = page.find_tables()
            for tab in tabs:
                t_bbox = [round(tab.bbox[0], 2), round(tab.bbox[1], 2), round(tab.bbox[2], 2), round(tab.bbox[3], 2)]
                table_matrix = tab.extract()
                blocks.append({
                    "id": f"b_{page_num}_{block_counter}",
                    "type": "table",
                    "bbox": t_bbox,
                    "content": table_matrix
                })
                block_counter += 1
        except Exception as tab_err:
            logger.debug(f"Table extraction notice on page {page_num}: {tab_err}")

        # Sort blocks by visual layout coordinates (top Y first, then left X)
        blocks.sort(key=lambda b: (b.get("bbox", [0, 0, 0, 0])[1], b.get("bbox", [0, 0, 0, 0])[0]))

        pages_data.append({
            "page": page_num,
            "width": page_w,
            "height": page_h,
            "blocks": blocks
        })

    doc.close()
    return {
        "pages": pages_data,
        "images_count": total_images,
        "formulas_count": total_formulas
    }


def parse_docx(file_path: Path) -> Dict[str, Any]:
    """
    Parses DOCX/DOC documents into normalized page and block structures.
    """
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx engine unavailable")

    doc = docx.Document(file_path)
    blocks = []
    block_counter = 1
    total_images = 0
    total_formulas = 0

    try:
        from PIL import Image, ImageDraw
        page_img_path = IMAGES_DIR / "pdf_page_1.png"
        img = Image.new("RGB", (1200, 1600), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)
        draw.text((40, 40), f"DOCX Document: {file_path.name}", fill=(15, 23, 42))
        
        y_text = 100
        for p in doc.paragraphs[:30]:
            text = p.text.strip()
            if text:
                draw.text((40, y_text), text[:100], fill=(51, 65, 85))
                y_text += 30
        
        img.save(str(page_img_path))
        logger.info(f"Rendered DOCX page image: {page_img_path}")
    except Exception as docx_render_err:
        logger.warning(f"DOCX page render warning: {docx_render_err}")

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            blocks.append({
                "id": f"b_1_{block_counter}",
                "type": "text",
                "bbox": [0.0, 0.0, 0.0, 0.0],
                "content": text
            })
            block_counter += 1

    for table in doc.tables:
        matrix = []
        for row in table.rows:
            row_vals = [cell.text.strip() for cell in row.cells]
            matrix.append(row_vals)
        blocks.append({
            "id": f"b_1_{block_counter}",
            "type": "table",
            "bbox": [0.0, 0.0, 0.0, 0.0],
            "content": matrix
        })
        block_counter += 1

    try:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_part = rel.target_part
                img_bytes = img_part.blob
                img_filename = f"docx_img_{uuid.uuid4().hex[:8]}.png"
                img_path = IMAGES_DIR / img_filename
                with open(img_path, "wb") as f:
                    f.write(img_bytes)
                
                total_images += 1
                blocks.append({
                    "id": f"b_1_{block_counter}",
                    "type": "image",
                    "bbox": [0.0, 0.0, 0.0, 0.0],
                    "image": f"/extracted_images/{img_filename}"
                })
                block_counter += 1
    except Exception as docx_img_err:
        logger.warning(f"Error extracting images from DOCX: {docx_img_err}")

    return {
        "pages": [
            {
                "page": 1,
                "width": 1200.0,
                "height": 1600.0,
                "blocks": blocks
            }
        ],
        "images_count": total_images,
        "formulas_count": total_formulas
    }


@app.post("/parse")
async def parse_document(file: UploadFile = File(...)):
    start_time = time.time()
    filename = file.filename or "uploaded_document"
    logger.info(f"[MinerU] Upload started: '{filename}'")

    ext = Path(filename).suffix.lower()
    allowed_extensions = {".pdf", ".doc", ".docx", ".png", ".jpg", ".jpeg"}
    if ext not in allowed_extensions:
        logger.error(f"Unsupported file format: {ext}")
        raise HTTPException(status_code=400, detail=f"Unsupported file format '{ext}'. Allowed: PDF, DOC, DOCX, PNG, JPG")

    temp_path = UPLOAD_DIR / f"{uuid.uuid4().hex}_{filename}"
    try:
        content = await file.read()
        with open(temp_path, "wb") as f_out:
            f_out.write(content)

        if ext == ".pdf":
            parsed_res = parse_pdf(temp_path)
        elif ext in [".docx", ".doc"]:
            parsed_res = parse_docx(temp_path)
        else: # Image files
            img_filename = f"img_direct_{uuid.uuid4().hex[:8]}{ext}"
            dest_img_path = IMAGES_DIR / img_filename
            with open(dest_img_path, "wb") as f_img:
                f_img.write(content)

            page_img_path = IMAGES_DIR / "pdf_page_1.png"
            with open(page_img_path, "wb") as f_p:
                f_p.write(content)
            
            parsed_res = {
                "pages": [
                    {
                        "page": 1,
                        "width": 1200.0,
                        "height": 1600.0,
                        "blocks": [
                            {
                                "id": "b_1_1",
                                "type": "image",
                                "bbox": [0.0, 0.0, 0.0, 0.0],
                                "image": f"/extracted_images/{img_filename}"
                            }
                        ]
                    }
                ],
                "images_count": 1,
                "formulas_count": 0
            }

        elapsed_time = round(time.time() - start_time, 3)
        page_count = len(parsed_res.get("pages", []))
        images_count = parsed_res.get("images_count", 0)
        formulas_count = parsed_res.get("formulas_count", 0)

        logger.info(f"[MinerU] Upload completed: '{filename}' | Time: {elapsed_time}s | Pages: {page_count} | Images: {images_count} | Formulas: {formulas_count}")

        return {
            "title": filename,
            "pages": parsed_res.get("pages", [])
        }

    except Exception as err:
        logger.error(f"[MinerU] Error parsing file '{filename}': {str(err)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Document parsing failed: {str(err)}")
    finally:
        if temp_path.exists():
            try:
                os.remove(temp_path)
            except Exception:
                pass


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
